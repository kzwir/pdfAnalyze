# -*- coding: utf-8 -*-

# 1. pobiera dane z engine
# 2. przetwarza tabelę
#   normalizacja kolumn
#   dodanie "nowa kolumna"
#   czyszczenie danych
# 3. waliduje
#   artefakty (1x1)
#   jakość danych (>30% pustych komórek)
# 4. generuje Word
# 5. zapisuje wynik

import os
import logging
import pandas as pd

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


logger = logging.getLogger(__name__)

# ============================================================
# DOCUMENT
# ============================================================

def create_document(template_path):
    try:
        document = Document(template_path)
    except Exception as e:
        logger.error(f"Błąd szablonu: {e}")
        document = Document()

    try:
        for table in document.tables:
            table._element.getparent().remove(table._element)
    except Exception as e:
        logger.warning(f"Nie udało się usunąć tabel z szablonu: {e}")

    return document


def save_document(document, output_path):
    try:
        os.makedirs(
            os.path.dirname(output_path) or ".",
            exist_ok=True
        )

        document.save(output_path)

        logger.info(
            f"Zapisano dokument: {output_path}"
        )

        return True

    except Exception as e:

        logger.error(
            f"Błąd zapisu dokumentu: {e}"
        )

        return False


# =========================
# WALIDACJA TABELI
# =========================
def is_valid_table(df):
    try:
        rows, cols = df.shape
        return rows > 1 and cols > 1
    except Exception:
        return False


# =========================
# NORMALIZACJA TABELI
# =========================
def normalize_table(table):
    """
    Ujednolica liczbę kolumn i uzupełnia brakujące wartości
    """
    try:
        max_cols = max(len(row) for row in table if row)

        normalized = []
        for row in table:
            r = row if row else []
            r = r + [""] * (max_cols - len(r))
            normalized.append(r)

        return normalized
    except Exception as e:
        logger.error(
            f"Błąd normalizacji tabeli: {e}"
        )
        return table



# =========================
# DODANIE NOWEJ KOLUMNY
# =========================
def add_new_column(df):
    try:
        df["nowa kolumna"] = ""

        if len(df) > 0:
            df.iat[0, df.shape[1] - 1] = "nowa kolumna"

    except Exception as e:
        logger.error(
            f"Błąd dodawania kolumny: {e}"
        )


    return df


# =========================
# JAKOŚĆ TABELI
# =========================
def check_quality(df):
    total_cells = df.shape[0] * df.shape[1]

    empty_cells = ((df == "") | (df.isna())).sum().sum()

    if total_cells == 0:
        return 0

    return empty_cells / total_cells


def calculate_quality(df):
    """
    Zwraca słownik metryk jakości.
    """

    try:

        rows, cols = df.shape

        total_cells = rows * cols

        if total_cells == 0:
            return None

        empty_cells = (
            (df == "") |
            (df.isna())
        ).sum().sum()

        empty_ratio = empty_cells / total_cells

        return {
            "rows": rows,
            "cols": cols,
            "total_cells": total_cells,
            "empty_cells": int(empty_cells),
            "empty_ratio": round(empty_ratio, 4)
        }

    except Exception as e:

        logger.error(
            f"Błąd wyliczania jakości: {e}"
        )

        return None


def validate_quality(df, table_counter):
    """
    Bardziej sensowna walidacja niż poprzedni próg 30%.
    """

    messages = []

    metrics = calculate_quality(df)

    if not metrics:
        return messages

    if metrics["rows"] <= 1:
        messages.append(
            f"Tabela {table_counter}: tylko jeden wiersz"
        )

    if metrics["cols"] <= 1:
        messages.append(
            f"Tabela {table_counter}: tylko jedna kolumna"
        )

    if metrics["empty_ratio"] > 0.95:

        messages.append(
            f"Tabela {table_counter}: "
            f"{metrics['empty_ratio']:.0%} komórek pustych "
            f"(prawdopodobny problem ekstrakcji)"
        )

    elif metrics["empty_ratio"] > 0.80:

        messages.append(
            f"Tabela {table_counter}: "
            f"{metrics['empty_ratio']:.0%} komórek pustych "
            f"(wymaga oceny użytkownika)"
        )

    return messages


# =========================
# USTAWIA CZARNE OBRAMOWANIE TABELI
# =========================
def set_table_borders(table):
    try:
        tbl = table._tbl

        tbl_pr = tbl.tblPr

        borders = OxmlElement("w:tblBorders")

        border_names = [
           "top",
           "left",
           "bottom",
           "right",
           "insideH",
           "insideV"
        ]

        for border_name in border_names:
           border = OxmlElement(f"w:{border_name}")
           border.set(qn("w:val"), "single")
           border.set(qn("w:sz"), "12")
           border.set(qn("w:space"), "0")
           border.set(qn("w:color"), "000000")
           borders.append(border)
        tbl_pr.append(borders)
    except Exception as e:
        logger.warning(
            f"Nie udało się ustawić obramowania tabeli: {e}"
        )

# ============================================================
# ENGINE FACTORY
# ============================================================

def create_engine(engine_name):

    try:

        if engine_name == "camelot":

            from engines.camelot_engine import (
                CamelotEngine
            )

            return CamelotEngine()

        if engine_name == "ocr":

            from engines.ocr_engine import (
                OCREngine
            )

            return OCREngine()

        from engines.pdfplumber_engine import (
            PdfPlumberEngine
        )

        return PdfPlumberEngine()

    except Exception as e:

        logger.error(
            f"Błąd inicjalizacji silnika: {e}"
        )

        raise


# =========================
# GŁÓWNA FUNKCJA
# =========================
def process_tables_to_word(
    tables,
    document,
    source_name="?"
):
    """
    Generuje zawartość Word na podstawie już
    wyekstrahowanych tabel.

    Nie wykonuje ekstrakcji PDF.
    """

    logger.info(
        f"Start TABLES -> Word: {source_name}"
    )

    validation_log = []

    table_counter = 1

    tables_found = False

    try:

        document.add_page_break()

        document.add_heading(
            f"Dokument źródłowy: {source_name}",
            level=1
        )

    except Exception as e:

        logger.warning(
            f"Błąd dodania nagłówka dokumentu: {e}"
        )

    valid_tables = 0

    for item in tables:

        try:

            table = item.get(
                "data",
                []
            )

            page = item.get(
                "page",
                "?"
            )

            if not table:
                continue

            tables_found = True

            raw_df = pd.DataFrame(
                table
            )

            if not is_valid_table(
                raw_df
            ):

                validation_log.append(
                    f"Tabela {table_counter}: "
                    f"pominięto artefakt"
                )

                table_counter += 1

                continue

            valid_tables += 1

            normalized = normalize_table(
                table
            )

            df = pd.DataFrame(
                normalized
            )

            df = add_new_column(df)

            validation_log.extend(
                validate_quality(
                    df,
                    table_counter
                )
            )

            document.add_paragraph(
                f"Tabela {table_counter} "
                f"(strona {page})"
            )

            table_doc = document.add_table(
                rows=df.shape[0],
                cols=df.shape[1]
            )

            set_table_borders(
                table_doc
            )

            table_doc.autofit = True

            try:

                table_doc.style = "zwir"

            except Exception:

                logger.debug(
                    "Styl 'zwir' nie istnieje."
                )

            for i in range(df.shape[0]):

                for j in range(df.shape[1]):

                    try:

                        value = df.iat[
                            i,
                            j
                        ]

                        table_doc.cell(
                            i,
                            j
                        ).text = (
                            str(value)
                            if value is not None
                            else ""
                        )

                    except Exception:

                        table_doc.cell(
                            i,
                            j
                        ).text = ""

            table_counter += 1

        except Exception as e:

            logger.error(
                f"Błąd tabeli "
                f"{table_counter}: {e}"
            )

            validation_log.append(
                f"Tabela {table_counter}: "
                f"błąd przetwarzania"
            )

            table_counter += 1

    if not tables_found:

        document.add_paragraph(
            "Nie wykryto tabel "
            "w dokumencie źródłowym"
        )

        validation_log.append(
            "Brak tabel"
        )

    try:

        document.add_page_break()

        document.add_heading(
            "Raport walidacji",
            level=1
        )

        if validation_log:

            for entry in validation_log:

                document.add_paragraph(
                    entry
                )

        else:

            document.add_paragraph(
                "Brak błędów"
            )

    except Exception as e:

        logger.warning(
            f"Błąd sekcji raportu: {e}"
        )

    result = {
        "tables": table_counter - 1,
        "tables_valid": valid_tables,
        "validation_log": validation_log
    }

    logger.info(
        f"Koniec TABLES -> Word: "
        f"{source_name}"
    )

    logger.info(
        f"Wynik: {result}"
    )

    return result

def process_pdf_to_word(
    pdf_path,
    document,
    engine_name="pdfplumber"
):
    """
    Funkcja kompatybilna z:
        glowny.py
        batch_pipeline.py
        tableimport.py
        parallel_pipeline.py
    """

    logger.info("pdf2word_module.py: process_pdf_to_word")
    logger.info(
        f"Start PDF -> Word: {pdf_path}"
    )

    validation_log = []

    table_counter = 1

    tables_found = False

    try:

        document.add_page_break()

        document.add_heading(
            f"Dokument źródłowy: "
            f"{os.path.basename(pdf_path)}",
            level=1
        )

    except Exception as e:

        logger.warning(
            f"Błąd dodania nagłówka dokumentu: {e}"
        )

    try:

        engine = create_engine(
            engine_name
        )

    except Exception:

        return {
            "tables": 0,
            "tables_valid": 0,
            "validation_log": [
                "Błąd inicjalizacji silnika"
            ]
        }

    try:

        tables = engine.extract(
            pdf_path
        )

    except Exception as e:

        logger.error(
            f"Błąd ekstrakcji: {e}"
        )

        tables = []

    valid_tables = 0

    for item in tables:

        try:

            table = item.get(
                "data",
                []
            )

            page = item.get(
                "page",
                "?"
            )

            if not table:
                continue

            tables_found = True

            raw_df = pd.DataFrame(
                table
            )

            if not is_valid_table(raw_df):

                validation_log.append(
                    f"Tabela {table_counter}: "
                    f"pominięto artefakt"
                )

                table_counter += 1

                continue

            valid_tables += 1

            normalized = normalize_table(
                table
            )

            df = pd.DataFrame(
                normalized
            )

            df = add_new_column(df)

            quality_messages = (
                validate_quality(
                    df,
                    table_counter
                )
            )

            validation_log.extend(
                quality_messages
            )

            document.add_paragraph(
                f"Tabela {table_counter} "
                f"(strona {page})"
            )

            table_doc = document.add_table(
                rows=df.shape[0],
                cols=df.shape[1]
            )

            set_table_borders(
                table_doc
            )

            table_doc.autofit = True

            try:
                table_doc.style = "zwir"
            except Exception:
                logger.debug(
                    "Styl 'zwir' nie istnieje."
                )

            for i in range(df.shape[0]):

                for j in range(df.shape[1]):

                    try:

                        value = df.iat[i, j]

                        cell = table_doc.cell(
                            i,
                            j
                        )

                        cell.text = (
                            str(value)
                            if value is not None
                            else ""
                        )

                    except Exception as e:

                        logger.warning(
                            f"Błąd komórki "
                            f"({i},{j}): {e}"
                        )

            table_counter += 1

        except Exception as e:

            logger.error(
                f"Błąd tabeli "
                f"{table_counter}: {e}"
            )

            validation_log.append(
                f"Tabela {table_counter}: "
                f"błąd przetwarzania"
            )

            table_counter += 1

    if not tables_found:

        document.add_paragraph(
            "Nie wykryto tabel "
            "w dokumencie źródłowym"
        )

        validation_log.append(
            "Brak tabel"
        )

    try:

        document.add_page_break()

        document.add_heading(
            "Raport walidacji",
            level=1
        )

        if validation_log:

            for entry in validation_log:

                document.add_paragraph(
                    entry
                )

        else:

            document.add_paragraph(
                "Brak błędów"
            )

    except Exception as e:

        logger.warning(
            f"Błąd sekcji raportu: {e}"
        )

    result = {
        "tables": table_counter - 1,
        "tables_valid": valid_tables,
        "validation_log": validation_log
    }

    logger.info(
        f"Koniec PDF -> Word: "
        f"{pdf_path}"
    )

    logger.info(
        f"Wynik: {result}"
    )

    return result
